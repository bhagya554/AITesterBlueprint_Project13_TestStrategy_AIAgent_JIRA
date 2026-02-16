"""
Export Service
Generates PDF and DOCX exports from markdown content
"""

import re
import io
from typing import Optional
from datetime import datetime


def markdown_to_html(content: str) -> str:
    """Convert markdown to basic HTML for PDF generation."""
    html = content
    
    # Escape HTML
    html = html.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    
    # Headers
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    
    # Bold and italic
    html = re.sub(r'\*\*\*(.+?)\*\*\*', r'<b><i>\1</i></b>', html)
    html = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', html)
    html = re.sub(r'\*(.+?)\*', r'<i>\1</i>', html)
    
    # Code
    html = re.sub(r'`(.+?)`', r'<code>\1</code>', html)
    
    # Line breaks and paragraphs
    html = html.replace('\n\n', '</p><p>')
    html = html.replace('\n', '<br>')
    
    # Wrap in body
    html = f"<body><p>{html}</p></body>"
    
    return html


def generate_pdf(
    content: str,
    title: str,
    project_name: Optional[str] = None,
    classification: str = "Confidential",
    company_name: Optional[str] = None
) -> bytes:
    """
    Generate PDF from markdown content using ReportLab.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    except ImportError:
        raise ImportError("reportlab not installed. Run: pip install reportlab")
    
    buffer = io.BytesIO()
    
    # Create document
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#1B3A5C')
    )
    
    heading1_style = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        spaceBefore=12,
        textColor=colors.HexColor('#1B3A5C')
    )
    
    heading2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=10,
        spaceBefore=10,
        textColor=colors.HexColor('#2E75B6')
    )
    
    heading3_style = ParagraphStyle(
        'CustomH3',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=8,
        textColor=colors.HexColor('#2E75B6')
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY
    )
    
    # Build content
    story = []
    
    # Cover page
    story.append(Spacer(1, 2*inch))
    
    if company_name:
        story.append(Paragraph(company_name, ParagraphStyle(
            'CompanyName',
            parent=styles['Normal'],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=30
        )))
    
    story.append(Paragraph("Test Strategy Document", title_style))
    story.append(Spacer(1, 0.5*inch))
    
    if project_name:
        story.append(Paragraph(f"Project: {project_name}", ParagraphStyle(
            'ProjectName',
            parent=styles['Normal'],
            fontSize=14,
            alignment=TA_CENTER,
            spaceAfter=20
        )))
    
    story.append(Paragraph(f"Title: {title}", ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontSize=12,
        alignment=TA_CENTER,
        spaceAfter=10
    )))
    
    story.append(Spacer(1, 1*inch))
    
    # Metadata table
    metadata_data = [
        ["Date:", datetime.now().strftime("%Y-%m-%d")],
        ["Classification:", classification],
        ["Version:", "1.0"],
    ]
    
    metadata_table = Table(metadata_data, colWidths=[1.5*inch, 3*inch])
    metadata_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
    ]))
    story.append(metadata_table)
    
    story.append(PageBreak())
    
    # Parse and add content
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 6))
            continue
        
        # Handle tables (simplified)
        if line.startswith('|') and line.endswith('|'):
            # This is a simplified table handling
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and not all(c.replace('-', '') == '' for c in cells):
                story.append(Paragraph(' | '.join(cells), body_style))
            continue
        
        # Headers
        if line.startswith('# '):
            story.append(Paragraph(line[2:], heading1_style))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], heading2_style))
        elif line.startswith('### '):
            story.append(Paragraph(line[4:], heading3_style))
        elif line.startswith('#### '):
            story.append(Paragraph(line[5:], heading3_style))
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            story.append(Paragraph(f"• {line[2:]}", body_style))
        elif re.match(r'^\d+\.', line):
            story.append(Paragraph(line, body_style))
        
        # Regular paragraph
        else:
            # Basic markdown formatting
            line = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', line)
            line = re.sub(r'\*(.+?)\*', r'<i>\1</i>', line)
            line = re.sub(r'`(.+?)`', r'<font face="Courier">\1</font>', line)
            story.append(Paragraph(line, body_style))
    
    # Build PDF
    doc.build(story)
    
    pdf_bytes = buffer.getvalue()
    buffer.close()
    
    return pdf_bytes


def generate_docx(
    content: str,
    title: str,
    project_name: Optional[str] = None,
    classification: str = "Confidential",
    company_name: Optional[str] = None
) -> bytes:
    """
    Generate DOCX from markdown content.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Cover page
    if company_name:
        p = doc.add_paragraph(company_name)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(16)
        run.font.bold = True
    
    p = doc.add_paragraph("Test Strategy Document")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    
    doc.add_paragraph()
    
    if project_name:
        p = doc.add_paragraph(f"Project: {project_name}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(14)
    
    p = doc.add_paragraph(f"Title: {title}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Metadata
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    table.rows[0].cells[0].text = "Date:"
    table.rows[0].cells[1].text = datetime.now().strftime("%Y-%m-%d")
    table.rows[1].cells[0].text = "Classification:"
    table.rows[1].cells[1].text = classification
    table.rows[2].cells[0].text = "Version:"
    table.rows[2].cells[1].text = "1.0"
    
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # Parse content and add to document
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Skip table separator lines
        if line.startswith('|') and '---' in line:
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Tables
        elif line.startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                # Add as paragraph for now (simplified)
                doc.add_paragraph(' | '.join(cells))
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        elif re.match(r'^\d+\.', line):
            p = doc.add_paragraph(line, style='List Number')
        
        # Regular paragraph
        else:
            p = doc.add_paragraph()
            
            # Simple markdown parsing for bold/italic
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
            
            for part in parts:
                run = p.add_run()
                
                if part.startswith('**') and part.endswith('**'):
                    run.text = part[2:-2]
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run.text = part[1:-1]
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run.text = part[1:-1]
                    run.font.name = 'Courier New'
                else:
                    run.text = part
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()
